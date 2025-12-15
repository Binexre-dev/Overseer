"""
Report Generation Module
Handles the generation of analysis reports in various formats
"""
import os
import json
from pathlib import Path
from typing import Dict, Any, List, Optional, Tuple
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


class ReportGenerator:
    """Generates analysis reports from collected data"""
    
    def __init__(self, analysis_dir: Path):
        self.analysis_dir = Path(analysis_dir)
        self.report_dir = self.analysis_dir / "reports"
        self.report_dir.mkdir(parents=True, exist_ok=True)
    
    def generate_report(
        self,
        analysis_results: Dict[str, Any],
        report_format: str = "docx"
    ) -> Path:
        """
        Generate a report from analysis results
        
        Args:
            analysis_results: Dictionary containing analysis data
            report_format: Format of report ('json', 'markdown', 'docx')
        
        Returns:
            Path to the generated report file
        """
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        binary_name = analysis_results.get("binary", {}).get("name", "unknown")
        
        if report_format == "json":
            return self._generate_json_report(analysis_results, timestamp, binary_name)
        elif report_format == "markdown":
            return self._generate_markdown_report(analysis_results, timestamp, binary_name)
        elif report_format == "docx":
            return self._generate_docx_report(analysis_results, timestamp, binary_name)
        else:
            raise ValueError(f"Unsupported report format: {report_format}")
    
    def generate_both_reports(self, analysis_results: Dict[str, Any]) -> Tuple[Path, Path]:
        """
        Generate both DOCX and JSON reports
        
        Args:
            analysis_results: Dictionary containing analysis data
        
        Returns:
            Tuple of (docx_path, json_path)
        """
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        binary_name = analysis_results.get("binary", {}).get("name", "unknown")
        
        docx_path = self._generate_docx_report(analysis_results, timestamp, binary_name)
        json_path = self._generate_json_report(analysis_results, timestamp, binary_name)
        
        return docx_path, json_path
    
    def _generate_json_report(
        self,
        results: Dict[str, Any],
        timestamp: str,
        binary_name: str
    ) -> Path:
        """Generate a JSON report"""
        report_path = self.report_dir / f"{binary_name}_{timestamp}.json"
        
        report_data = {
            "binary": binary_name,
            "timestamp": timestamp,
            "analysis_results": results
        }
        
        with open(report_path, 'w', encoding='utf-8') as f:
            json.dump(report_data, f, indent=4)
        
        return report_path
    
    def _generate_docx_report(
        self,
        results: Dict[str, Any],
        timestamp: str,
        binary_name: str
    ) -> Path:
        """Generate a DOCX report"""
        report_path = self.report_dir / f"{binary_name}_{timestamp}.docx"
        
        doc = Document()
        
        # Title
        title = doc.add_heading('Malware Analysis Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Summary Section
        doc.add_heading('Summary', 1)
        summary_table = doc.add_table(rows=3, cols=2)
        summary_table.style = 'Light Grid Accent 1'
        
        summary_data = [
            ('Binary:', binary_name),
            ('Analysis Date:', timestamp),
            ('Binary Path:', results.get('binary', {}).get('path', 'N/A'))
        ]
        
        for i, (label, value) in enumerate(summary_data):
            row = summary_table.rows[i]
            row.cells[0].text = label
            row.cells[0].paragraphs[0].runs[0].bold = True
            row.cells[1].text = str(value)
        
        # Static Analysis Results
        doc.add_heading('Static Analysis Results', 1)
        static_results = results.get('static_results', {})
        if static_results:
            for tool_name, result in static_results.items():
                doc.add_heading(tool_name, 2)
                doc.add_paragraph(json.dumps(result, indent=2))
        else:
            doc.add_paragraph('No static analysis results available.')
        
        # Dynamic Analysis Results
        doc.add_heading('Dynamic Analysis Results', 1)
        dynamic_results = results.get('dynamic_results', {})
        if dynamic_results:
            for tool_name, result in dynamic_results.items():
                doc.add_heading(tool_name, 2)
                doc.add_paragraph(json.dumps(result, indent=2))
        else:
            doc.add_paragraph('No dynamic analysis results available.')
        
        # Procmon Results
        doc.add_heading('Process Monitor Results', 1)
        procmon_results = results.get('procmon_results', {})
        if procmon_results:
            doc.add_paragraph(json.dumps(procmon_results, indent=2))
        else:
            doc.add_paragraph('No Process Monitor data available.')
        
        doc.save(str(report_path))
        return report_path
    
    def _generate_markdown_report(
        self,
        results: Dict[str, Any],
        timestamp: str,
        binary_name: str
    ) -> Path:
        """Generate a Markdown report"""
        report_path = self.report_dir / f"{binary_name}_{timestamp}.md"
        
        md_content = f"""# Malware Analysis Report

## Summary
- **Binary:** {binary_name}
- **Analysis Date:** {timestamp}
- **Binary Path:** {results.get('binary', {}).get('path', 'N/A')}

## Static Analysis Results
{self._format_tool_results_markdown(results.get('static_results', {}))}

## Dynamic Analysis Results
{self._format_tool_results_markdown(results.get('dynamic_results', {}))}

## Procmon Results
{self._format_procmon_results_markdown(results.get('procmon_results', {}))}
"""
        
        with open(report_path, 'w', encoding='utf-8') as f:
            f.write(md_content)
        
        return report_path
    
    def _format_tool_results_markdown(self, tool_results: Dict[str, Any]) -> str:
        """Format tool results for Markdown"""
        if not tool_results:
            return "No results available\n"
        
        md = ""
        for tool_name, result in tool_results.items():
            md += f"\n### {tool_name}\n\n```json\n{json.dumps(result, indent=2)}\n```\n"
        return md
    
    def _format_procmon_results_markdown(self, procmon_results: Dict[str, Any]) -> str:
        """Format Procmon results for Markdown"""
        if not procmon_results:
            return "No Procmon data available\n"
        
        return f"\n```json\n{json.dumps(procmon_results, indent=2)}\n```\n"


def generate_analysis_report(
    analysis_results: Dict[str, Any],
    output_dir: Path,
    report_format: str = "both"
) -> Tuple[Path, Path] | Path:
    """
    Convenience function to generate a report
    
    Args:
        analysis_results: Dictionary containing analysis data
        output_dir: Directory to save the report
        report_format: Format of report ('json', 'markdown', 'docx', 'both')
                      'both' generates both DOCX and JSON
    
    Returns:
        Path to the generated report file, or tuple of (docx_path, json_path) if format='both'
    """
    generator = ReportGenerator(output_dir)
    
    if report_format == "both":
        return generator.generate_both_reports(analysis_results)
    else:
        return generator.generate_report(analysis_results, report_format)
